#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import os
import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Helper: prefix Persian paragraphs with RLM to reduce bidi issues
RLM = '\u200F'
LRM = '\u200E'

IGNORED_DIRS = {'.git', '__pycache__', '.venv', 'venv', 'node_modules'}
MAX_LINES = 120


def add_persian_paragraph(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run(RLM + text)
    font = run.font
    font.name = 'Tahoma'
    font.size = Pt(size)
    run.bold = bold
    return p


def add_code_block(doc, code, size=9):
    # code should remain LTR and monospaced
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run(code)
    font = run.font
    font.name = 'Consolas'
    font.size = Pt(size)
    return p


def summarize_repo(root_dir):
    doc = Document()

    title = 'خلاصهٔ کد پروژه MiniMony'
    add_persian_paragraph(doc, title, size=18, bold=True)

    date_str = datetime.datetime.now().strftime('%Y-%m-%d %H:%M')
    add_persian_paragraph(doc, f'تاریخ تولید سند: {date_str}', size=10)

    intro = ('این سند به‌صورت خودکار از محتوای مخزن ساخته شده است. ' 
             'فایل‌ها و بخشی از محتوای آن‌ها در ادامه آمده است. ' 
             'برای جلوگیری از به‌هم‌ریختگی جهت (LTR/RTL)، جمله‌های فارسی در سمت راست قرار گرفته‌اند و بخش‌های کد در سمت چپ.')
    add_persian_paragraph(doc, intro)

    # Walk repository
    for dirpath, dirnames, filenames in os.walk(root_dir):
        # prune ignored dirs
        parts = dirpath.split(os.sep)
        if any(p in IGNORED_DIRS for p in parts):
            continue

        # relative path
        rel_dir = os.path.relpath(dirpath, root_dir)
        if rel_dir == '.':
            rel_dir = '/'
        add_persian_paragraph(doc, f'پوشه: {rel_dir}', bold=True)

        for fname in sorted(filenames):
            # skip binary-like large files
            if fname.endswith(('.pyc', '.exe', '.bin', '.png', '.jpg', '.jpeg', '.zip', '.gz')):
                continue
            filepath = os.path.join(dirpath, fname)
            rel_path = os.path.relpath(filepath, root_dir)

            add_persian_paragraph(doc, f'فایل: {rel_path}', size=11, bold=True)

            try:
                with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                    lines = f.readlines()
            except Exception as e:
                add_persian_paragraph(doc, f'خطا در خواندن فایل: {e}')
                continue

            if not lines:
                add_persian_paragraph(doc, 'فایل خالی است.')
                continue

            snippet = ''.join(lines[:MAX_LINES])
            # Ensure code snippet uses LTR markers to preserve formatting
            snippet = LRM + snippet
            add_code_block(doc, snippet)

    return doc


if __name__ == '__main__':
    ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
    out_path = os.path.join(ROOT, 'MiniMony_summary_fa.docx')
    doc = summarize_repo(ROOT)
    doc.save(out_path)
    print('Saved:', out_path)
